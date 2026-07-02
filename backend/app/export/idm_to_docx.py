from __future__ import annotations

from io import BytesIO

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, RGBColor

from app.models.idm import Block, DocumentModel, Run

ALIGNMENT = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}


def idm_to_docx(document: DocumentModel) -> bytes:
    doc = Document()
    for page in document.pages:
        for block in page.blocks:
            _add_block(doc, block)
    stream = BytesIO()
    doc.save(stream)
    return stream.getvalue()


def _add_block(doc: Document, block: Block) -> None:
    if block.type == "heading":
        paragraph = doc.add_heading(level=int(block.attrs.get("level", 1)))
        _add_runs(paragraph, block.runs or [])
    elif block.type in {"paragraph", "signatureLine", "clauseNumber"}:
        paragraph = doc.add_paragraph()
        if block.type == "paragraph":
            paragraph.alignment = ALIGNMENT.get(block.attrs.get("alignment", "left"), WD_ALIGN_PARAGRAPH.LEFT)
            paragraph.paragraph_format.left_indent = Inches(0.25 * int(block.attrs.get("indentLevel", 0)))
        _add_runs(paragraph, block.runs or [])
    elif block.type == "pageBreak":
        doc.add_page_break()
    elif block.type == "list":
        for child in block.children or []:
            paragraph = doc.add_paragraph(style="List Bullet" if block.attrs.get("listType") == "bullet" else "List Number")
            paragraph.paragraph_format.left_indent = Inches(0.25 * int(child.attrs.get("indentLevel", 0)))
            _add_runs(paragraph, child.runs or [])
    elif block.type == "table":
        _add_table(doc, block)


def _add_table(doc: Document, block: Block) -> None:
    rows = block.children or []
    row_count = max(1, int(block.attrs.get("rowCount", len(rows))))
    col_count = max(1, int(block.attrs.get("colCount", 1)))
    table = doc.add_table(rows=row_count, cols=col_count)
    table.style = "Table Grid"
    for row_block in rows:
        for cell_block in row_block.children or []:
            row = int(cell_block.attrs.get("rowIndex", 0))
            col = int(cell_block.attrs.get("colIndex", 0))
            row_span = int(cell_block.attrs.get("rowSpan", 1))
            col_span = int(cell_block.attrs.get("colSpan", 1))
            cell = table.cell(row, col)
            target = table.cell(row + row_span - 1, col + col_span - 1)
            if target is not cell:
                cell = cell.merge(target)
            paragraph = cell.paragraphs[0]
            _add_runs(paragraph, cell_block.runs or [])


def _add_runs(paragraph, runs: list[Run]) -> None:
    for source in runs:
        run = paragraph.add_run(source.text)
        run.bold = source.marks.bold
        run.italic = source.marks.italic
        run.underline = source.marks.underline
        if source.marks.color:
            run.font.color.rgb = RGBColor.from_string(source.marks.color.lstrip("#"))
